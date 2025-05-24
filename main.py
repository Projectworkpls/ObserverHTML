import streamlit as st
import requests
import base64
import json
from supabase import create_client
from datetime import datetime, timedelta
import re
import time
import google.generativeai as genai
import docx
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import os
import pathlib
import uuid
import logging
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import calendar

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set page config
st.set_page_config(
    page_title="Learning Observer",
    layout="wide",
    page_icon="📝"
)


# Initialize Supabase client with enhanced error handling
@st.cache_resource
def init_supabase():
    try:
        SUPABASE_URL = st.secrets.get("SUPABASE_URL")
        SUPABASE_KEY = st.secrets.get("SUPABASE_KEY")

        if not SUPABASE_URL or not SUPABASE_KEY:
            raise ValueError("Supabase URL or KEY not found in secrets")

        client = create_client(SUPABASE_URL, SUPABASE_KEY)

        # Test connection
        try:
            test = client.table('users').select("count", count="exact").execute()
            logger.info(f"Supabase connected successfully. Found {test.count} users.")
        except Exception as test_error:
            logger.error(f"Supabase connection test failed: {str(test_error)}")
            st.error("Failed to connect to database. Please check your connection settings.")
            raise

        return client

    except Exception as e:
        logger.error(f"Supabase initialization failed: {str(e)}")
        st.error("Database initialization failed. Please check your configuration.")
        raise


supabase = init_supabase()


def upload_file_to_storage(file_data, file_name, file_type):
    """Upload a file to Supabase Storage and return the URL"""
    try:
        # Create a unique file name to avoid collisions
        unique_filename = f"{uuid.uuid4()}_{file_name}"

        # Determine the appropriate bucket based on file type
        bucket_name = "audio-files" if "audio" in file_type else "image-files"

        # Upload the file to Supabase Storage
        response = supabase.storage.from_(bucket_name).upload(
            unique_filename,
            file_data,
            file_options={"content-type": file_type}
        )

        # Get the public URL for the file
        file_url = supabase.storage.from_(bucket_name).get_public_url(unique_filename)

        return file_url
    except Exception as e:
        st.error(f"Error uploading file to storage: {str(e)}")
        return None


# Configure Google AI API
genai.configure(api_key=st.secrets.get("GOOGLE_API_KEY"))

# Set up AssemblyAI API key
assemblyai_key = st.secrets.get("ASSEMBLYAI_API_KEY", "")


class ObservationExtractor:
    def __init__(self):
        self.ocr_api_key = st.secrets.get("OCR_API_KEY")
        self.groq_api_key = st.secrets.get("GROQ_API_KEY")
        self.gemini_api_key = st.secrets.get("GOOGLE_API_KEY")

    def image_to_base64(self, image_file):
        """Convert image file to base64 string"""
        return base64.b64encode(image_file.read()).decode('utf-8')

    def extract_text_with_ocr(self, image_file):
        """Extract text from image using OCR.space API"""
        try:
            # Get file extension
            file_type = image_file.name.split('.')[-1].lower()
            if file_type == 'jpeg':
                file_type = 'jpg'

            # Convert image to base64
            base64_image = self.image_to_base64(image_file)
            base64_image_with_prefix = f"data:image/{file_type};base64,{base64_image}"

            # Prepare request payload
            payload = {
                'apikey': self.ocr_api_key,
                'language': 'eng',
                'isOverlayRequired': False,
                'iscreatesearchablepdf': False,
                'issearchablepdfhidetextlayer': False,
                'OCREngine': 2,  # Better for handwriting
                'detectOrientation': True,
                'scale': True,
                'base64Image': base64_image_with_prefix
            }

            # Send request to OCR API
            response = requests.post(
                'https://api.ocr.space/parse/image',
                data=payload,
                headers={'apikey': self.ocr_api_key}
            )

            response.raise_for_status()
            data = response.json()

            # Process response
            if not data.get('ParsedResults') or len(data['ParsedResults']) == 0:
                error_msg = data.get('ErrorMessage', 'No parsed results returned')
                raise Exception(f"OCR Error: {error_msg}")

            parsed_result = data['ParsedResults'][0]
            if parsed_result.get('ErrorMessage'):
                raise Exception(f"OCR Error: {parsed_result['ErrorMessage']}")

            extracted_text = parsed_result['ParsedText']

            if not extracted_text or not extracted_text.strip():
                raise Exception("No text was detected in the image")

            return extracted_text

        except Exception as e:
            st.error(f"OCR Error: {str(e)}")
            raise

    def process_with_groq(self, extracted_text):
        """Process extracted text with Groq AI"""
        try:
            # Original detailed prompt
            system_prompt = """You are an AI assistant for a learning observation system. Extract and structure information from the provided observation sheet text.
The observation sheets typically have the following structure:
- Title (usually "The Observer")
- Student information (Name, Roll Number/ID)
- Date and Time information
- Core Observation Section with time slots
- Teaching content for each time slot
- Learning details (what was learned, tools used, etc.)

Format your response as JSON with the following structure:
{
  "studentName": "Student's name if available, otherwise use the title of the observation",
  "studentId": "Student ID or Roll Number",
  "className": "Class name or subject being taught",
  "date": "Date of observation",
  "observations": "Detailed description of what was learned",
  "strengths": ["List of strengths observed in the student"],
  "areasOfDevelopment": ["List of areas where the student needs improvement"],
  "recommendations": ["List of recommended actions for improvement"]
}

For observations, provide full detailed descriptions like:
"The student learned how to make maggi from their mom through in-person mode, including all steps from boiling water to adding spices"

Be creative in extracting information based on context."""

            # Send request to Groq API
            response = requests.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {self.groq_api_key}',
                    'Content-Type': 'application/json'
                },
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": f"Extract and structure the following text from an observation sheet: {extracted_text}"
                        }
                    ],
                    "temperature": 0.2,
                    "response_format": {"type": "json_object"}
                }
            )

            response.raise_for_status()
            data = response.json()

            # Extract the JSON content
            ai_response = data['choices'][0]['message']['content']
            return json.loads(ai_response)

        except Exception as e:
            st.error(f"Groq API Error: {str(e)}")
            raise

    def transcribe_with_assemblyai(self, audio_file):
        """Transcribe audio using AssemblyAI API"""
        if not assemblyai_key:
            return "Error: AssemblyAI API key is not configured. Please add it to your secrets."

        # Set up the API headers
        headers = {
            "authorization": assemblyai_key,
            "content-type": "application/json"
        }

        # Upload the audio file
        try:
            st.write("Uploading audio file...")
            upload_response = requests.post(
                "https://api.assemblyai.com/v2/upload",
                headers={"authorization": assemblyai_key},
                data=audio_file.getvalue()
            )

            if upload_response.status_code != 200:
                return f"Error uploading audio: {upload_response.text}"

            upload_url = upload_response.json()["upload_url"]

            # Request transcription
            st.write("Processing transcription...")
            transcript_request = {
                "audio_url": upload_url,
                "language_code": "en"
            }

            transcript_response = requests.post(
                "https://api.assemblyai.com/v2/transcript",
                json=transcript_request,
                headers=headers
            )

            if transcript_response.status_code != 200:
                return f"Error requesting transcription: {transcript_response.text}"

            transcript_id = transcript_response.json()["id"]

            # Poll for completion
            status = "processing"
            progress_bar = st.progress(0)
            while status != "completed" and status != "error":
                polling_response = requests.get(
                    f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
                    headers=headers
                )

                if polling_response.status_code != 200:
                    return f"Error checking transcription status: {polling_response.text}"

                polling_data = polling_response.json()
                status = polling_data["status"]

                if status == "completed":
                    progress_bar.progress(100)
                    return polling_data["text"]
                elif status == "error":
                    return f"Transcription error: {polling_data.get('error', 'Unknown error')}"

                # Update progress
                progress = polling_data.get("percent_done", 0)
                if progress:
                    progress_bar.progress(progress / 100.0)
                time.sleep(2)

            return "Error: Transcription timed out or failed."
        except Exception as e:
            return f"Error during transcription: {str(e)}"

    def generate_report_from_text(self, text_content, user_info):
        """Generate a structured report from text using Google Gemini"""
        prompt = f"""
        Based on this text from a student observation, create a detailed observer report following the new Daily Growth Report format.

        TEXT CONTENT:
        {text_content}

        FORMAT REQUIREMENTS:

        🧾 Daily Growth Report Format for Parents

        🧒 Child's Name: {user_info['student_name']}
        📅 Date: [{user_info['session_date']}]
        🎯 Theme of the Day: [Extract from text]
        🌱 Curiosity Seed Explored: [Extract from text]

        📊 Growth Metrics & Observations
        Growth Area | Rating | Observation Summary
        🧠 Intellectual | [✅ Excellent/✅ Good/⚠️ Fair/❌ Needs Work] | [Brief summary]
        😊 Emotional | [✅ Excellent/✅ Good/⚠️ Fair/❌ Needs Work] | [Brief summary]
        🤝 Social | [✅ Excellent/✅ Good/⚠️ Fair/❌ Needs Work] | [Brief summary]
        🎨 Creativity | [✅ Excellent/✅ Good/⚠️ Fair/❌ Needs Work] | [Brief summary]
        🏃 Physical | [✅ Excellent/✅ Good/⚠️ Fair/❌ Needs Work] | [Brief summary]
        🧭 Character/Values | [✅ Excellent/✅ Good/⚠️ Fair/❌ Needs Work] | [Brief summary]
        🚀 Planning/Independence | [✅ Excellent/✅ Good/⚠️ Fair/❌ Needs Work] | [Brief summary]

        🌈 Curiosity Response Index: [1-10] / 10
        [Brief explanation of child's engagement with the curiosity seed]

        🧠 Overall Growth Score:
        [🔵 Balanced Growth/🟡 Moderate Growth/🔴 Limited Growth] – [X/7] Areas Active
        [Brief recommendation]

        📣 Parent Note:
        [Comprehensive summary for parents with actionable insights]

        🟢 Legend:
        ✅ Excellent: Clear growth with evidence
        ⚠️ Fair: Some engagement, needs encouragement
        ❌ Needs Work: Area not activated today

        🔵 Overall Score:
        🟢 Good (6-7 active areas)
        🟡 Moderate (3-5 areas)
        🔴 Needs Encouragement (1-2 areas)

        Use the exact section titles, emojis and format as above. For items that cannot be determined from the text, make reasonable inferences based on the available information.
        """

        try:
            # Configure the model - using Gemini Pro for most comprehensive responses
            model = genai.GenerativeModel('gemini-2.0-flash')

            # Generate content with Gemini
            response = model.generate_content([
                {"role": "user", "parts": [{"text": prompt}]}
            ])

            # Extract the content from the response
            report_content = response.text

            return report_content
        except Exception as e:
            return f"Error generating report: {str(e)}"

    def create_word_document(self, report_content):
        """Create a Word document from the report content with proper formatting"""
        doc = docx.Document()

        # Add title
        title = doc.add_heading('Daily Growth Report', 0)

        # Clean up markdown formatting
        report_content = report_content.replace('**', '')

        # Add report content, parsing the sections
        lines = report_content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Header information (Child's Name, Date, etc.)
            if line.startswith(('🧒', '📅', '🎯', '🌱')):
                p = doc.add_paragraph()
                p.add_run(line).bold = True

            # Growth Metrics header
            elif line.startswith('📊'):
                doc.add_heading(line, level=1)

            # Growth area rows
            elif line.startswith(('🧠', '😊', '🤝', '🎨', '🏃', '🧭', '🚀')):
                p = doc.add_paragraph()
                p.add_run(line).bold = True

            # Curiosity Response Index
            elif line.startswith('🌈'):
                doc.add_heading(line, level=2)

            # Overall Growth Score
            elif line.startswith('🧠 Overall'):
                doc.add_heading(line, level=2)

            # Parent Note
            elif line.startswith('📣'):
                doc.add_heading(line, level=2)

            # Legend
            elif line.startswith('🟢 Legend'):
                doc.add_heading(line, level=3)

            # Legend items
            elif line.startswith(('✅', '⚠️', '❌', '🔵', '🟢', '🟡', '🔴')):
                p = doc.add_paragraph(line, style='List Bullet')

            # Regular paragraph
            else:
                doc.add_paragraph(line)

        # Save to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        return docx_bytes

    def send_email(self, recipient_email, subject, message):
        """Send email with the observation report"""
        sender_email = "parth.workforai@gmail.com"
        sender_password = st.secrets.get("EMAIL_PASSWORD")  # Add this to your secrets.toml

        if not sender_password:
            return False, "Email password not configured in secrets"

        smtp_server = "smtp.gmail.com"
        smtp_port = 587

        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = recipient_email
        msg["Subject"] = subject
        msg.attach(MIMEText(message, "html"))

        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            return True, f"Email sent to {recipient_email}"
        except smtplib.SMTPAuthenticationError:
            return False, "Error: Authentication failed. Check your email and password."
        except smtplib.SMTPException as e:
            return False, f"Error: Failed to send email. {e}"
        except Exception as e:
            return False, f"Error: {str(e)}"
        finally:
            try:
                server.quit()
            except:
                pass


class MonthlyReportGenerator:
    def __init__(self, supabase_client):
        self.supabase = supabase_client

    def get_month_data(self, child_id, year, month):
        """Fetch all observations for a specific child in a given month"""
        # Convert month/year to date range
        start_date = f"{year}-{month:02d}-01"
        if month == 12:
            end_date = f"{year + 1}-01-01"
        else:
            end_date = f"{year}-{month + 1:02d}-01"

        try:
            # Query observations table for the specified date range
            response = self.supabase.table('observations').select("*") \
                .eq("student_id", child_id) \
                .gte("date", start_date) \
                .lt("date", end_date) \
                .execute()

            return response.data
        except Exception as e:
            st.error(f"Error fetching monthly data: {str(e)}")
            return []

    def get_goal_progress(self, child_id, year, month):
        """Get goal progress data for the specified month"""
        # Convert month/year to date range
        start_date = f"{year}-{month:02d}-01"
        if month == 12:
            end_date = f"{year + 1}-01-01"
        else:
            end_date = f"{year}-{month + 1:02d}-01"

        try:
            # Get goals for this child
            goals_response = self.supabase.table('goals').select("*") \
                .eq("child_id", child_id) \
                .execute()

            goals = goals_response.data

            # Get alignments for these goals within the time period
            goal_progress = []

            for goal in goals:
                # Get alignments for this goal
                alignments_response = self.supabase.table('goal_alignments').select("*") \
                    .eq("goal_id", goal['id']) \
                    .execute()

                alignments = alignments_response.data

                # Filter alignments by report date
                relevant_alignments = []
                for alignment in alignments:
                    report_response = self.supabase.table('observations').select("date") \
                        .eq("id", alignment['report_id']) \
                        .execute()

                    if report_response.data:
                        report_date = report_response.data[0]['date']
                        if start_date <= report_date < end_date:
                            relevant_alignments.append(alignment)

                # Calculate progress metrics
                if relevant_alignments:
                    avg_score = sum(a['alignment_score'] for a in relevant_alignments) / len(relevant_alignments)
                    progress_trend = [a['alignment_score'] for a in relevant_alignments]

                    goal_progress.append({
                        'goal_text': goal['goal_text'],
                        'avg_score': avg_score,
                        'progress_trend': progress_trend,
                        'num_observations': len(relevant_alignments)
                    })

            return goal_progress
        except Exception as e:
            st.error(f"Error fetching goal progress: {str(e)}")
            return []

    def get_strength_areas(self, observations):
        """Extract and count strength areas from observations"""
        strength_counts = {}

        for obs in observations:
            if obs.get('strengths'):
                try:
                    strengths = json.loads(obs['strengths']) if isinstance(obs['strengths'], str) else obs['strengths']
                    for strength in strengths:
                        strength_counts[strength] = strength_counts.get(strength, 0) + 1
                except:
                    pass

        # Sort by frequency
        return dict(sorted(strength_counts.items(), key=lambda x: x[1], reverse=True))

    def get_development_areas(self, observations):
        """Extract and count development areas from observations"""
        development_counts = {}

        for obs in observations:
            if obs.get('areas_of_development'):
                try:
                    areas = json.loads(obs['areas_of_development']) if isinstance(obs['areas_of_development'], str) else \
                        obs['areas_of_development']
                    for area in areas:
                        development_counts[area] = development_counts.get(area, 0) + 1
                except:
                    pass

        # Sort by frequency
        return dict(sorted(development_counts.items(), key=lambda x: x[1], reverse=True))

    def generate_observation_frequency_chart(self, observations):
        """Generate a chart showing the frequency of observations by date"""
        # Extract dates and count observations per date
        date_counts = {}

        for obs in observations:
            date = obs.get('date', '')
            if date:
                date_counts[date] = date_counts.get(date, 0) + 1

        if not date_counts:
            return None

        # Create dataframe for plotting
        df = pd.DataFrame([
            {"date": date, "count": count}
            for date, count in date_counts.items()
        ])

        # Sort by date
        df['date'] = pd.to_datetime(df['date'])
        df = df.sort_values('date')

        # Create the figure
        fig = px.bar(
            df,
            x='date',
            y='count',
            title='Observation Frequency by Date',
            labels={'date': 'Date', 'count': 'Number of Observations'}
        )

        return fig

    def generate_strengths_chart(self, strength_counts):
        """Generate a chart showing the frequency of different strengths"""
        if not strength_counts:
            return None

        # Take top 10 strengths
        top_strengths = dict(list(strength_counts.items())[:10])

        # Create dataframe for plotting
        df = pd.DataFrame([
            {"strength": strength, "count": count}
            for strength, count in top_strengths.items()
        ])

        # Create the figure
        fig = px.bar(
            df,
            x='count',
            y='strength',
            title='Top Strengths Observed',
            labels={'strength': 'Strength', 'count': 'Frequency'},
            orientation='h'
        )

        return fig

    def generate_development_areas_chart(self, development_counts):
        """Generate a chart showing the frequency of different development areas"""
        if not development_counts:
            return None

        # Take top 10 development areas
        top_areas = dict(list(development_counts.items())[:10])

        # Create dataframe for plotting
        df = pd.DataFrame([
            {"area": area, "count": count}
            for area, count in top_areas.items()
        ])

        # Create the figure
        fig = px.bar(
            df,
            x='count',
            y='area',
            title='Areas for Development',
            labels={'area': 'Development Area', 'count': 'Frequency'},
            orientation='h'
        )

        return fig

    def generate_goal_progress_chart(self, goal_progress):
        """Generate a chart showing progress on goals"""
        if not goal_progress:
            return None

        # Create figure with subplots
        fig = make_subplots(rows=len(goal_progress), cols=1,
                            subplot_titles=[g['goal_text'][:50] + '...' for g in goal_progress],
                            vertical_spacing=0.1)

        for i, goal in enumerate(goal_progress):
            # Add bar for average score
            fig.add_trace(
                go.Bar(
                    x=[goal['avg_score']],
                    y=['Average Score'],
                    orientation='h',
                    name=f"Goal {i + 1}",
                    showlegend=False
                ),
                row=i + 1, col=1
            )

            # Add a reference line for goal target (10)
            fig.add_shape(
                type="line",
                x0=10, y0=-0.5,
                x1=10, y1=0.5,
                line=dict(color="green", width=2, dash="dash"),
                row=i + 1, col=1
            )

        # Update layout
        fig.update_layout(
            title_text="Goal Progress",
            height=200 * len(goal_progress),
            margin=dict(l=0, r=0, t=50, b=0)
        )

        return fig

    def generate_monthly_summary(self, observations, goal_progress):
        """Generate a text summary of the monthly progress"""
        if not observations:
            return "No observations recorded this month."

        num_observations = len(observations)
        num_goals_with_progress = len(goal_progress)

        # Calculate overall progress
        if goal_progress:
            avg_goal_score = sum(g['avg_score'] for g in goal_progress) / len(goal_progress)
            highest_goal = max(goal_progress, key=lambda x: x['avg_score'])
            lowest_goal = min(goal_progress, key=lambda x: x['avg_score'])
        else:
            avg_goal_score = 0
            highest_goal = None
            lowest_goal = None

        # Create summary
        summary = f"""
        ### Monthly Progress Summary

        **Total Observations:** {num_observations}

        **Goals Tracked:** {num_goals_with_progress}

        **Average Goal Progress:** {avg_goal_score:.1f}/10
        """

        if highest_goal:
            summary += f"""
            **Strongest Goal Area:** {highest_goal['goal_text'][:50]}... (Score: {highest_goal['avg_score']:.1f}/10)

            **Goal Needing Most Support:** {lowest_goal['goal_text'][:50]}... (Score: {lowest_goal['avg_score']:.1f}/10)
            """

        return summary


def admin_dashboard():
    st.title("Admin Dashboard")

    tabs = st.tabs(
        ["User Management", "Parent-Child Mappings", "Observer-Child Mappings", "Activity Logs", "Report Processing"])

    with tabs[0]:  # User Management
        st.subheader("User Management")
        try:
            users = supabase.table('users').select("*").execute().data
            if users:
                st.write("Registered Users:")
                for user in users:
                    with st.expander(f"{user.get('name', 'N/A')} ({user['email']}) - {user['role']}"):
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            # Display user details
                            st.write(f"ID: {user['id']}")
                            st.write(f"Role: {user['role']}")
                            st.write(f"Created: {user.get('created_at', 'N/A')}")

                            # Special fields based on role
                            if user['role'] == 'Parent':
                                child_id = user.get('child_id')
                                if child_id:
                                    child = supabase.table('children').select("*").eq('id', child_id).execute().data
                                    if child:
                                        st.write(f"Assigned Child: {child[0].get('id', child_id)}")
                                    else:
                                        st.write(f"Assigned Child ID: {child_id} (not found)")
                                else:
                                    st.write("No child assigned")

                        with col2:
                            if st.button("Delete", key=f"delete_{user['id']}"):
                                supabase.table('users').delete().eq('id', user['id']).execute()
                                st.rerun()
            else:
                st.info("No users found")
        except Exception as e:
            st.error(f"Database error: {str(e)}")

    with tabs[1]:  # Parent-Child Mappings
        st.subheader("Parent-Child Relationships")

        # Add new section for bulk child upload
        with st.expander("Bulk Add Children (CSV)"):
            st.info(
                "Upload a CSV file with one column 'name' for child names (optional columns: 'birth_date', 'grade')")
            child_upload = st.file_uploader("Upload Children CSV", type=["csv"], key="child_upload")

            if child_upload:
                try:
                    df = pd.read_csv(child_upload)

                    if 'name' not in df.columns:
                        st.error("CSV must contain a 'name' column")
                    else:
                        st.write("Preview of children to be added:")
                        st.dataframe(df.head())

                        if st.button("Add Children"):
                            children_data = []
                            for _, row in df.iterrows():
                                child_data = {
                                    "name": row['name'],
                                    "birth_date": row.get('birth_date', None),
                                    "grade": row.get('grade', None)
                                }
                                children_data.append(child_data)

                            # Insert in batches to avoid timeout
                            batch_size = 50
                            for i in range(0, len(children_data), batch_size):
                                batch = children_data[i:i + batch_size]
                                supabase.table('children').insert(batch).execute()

                            st.success(f"Successfully added {len(children_data)} children!")
                            st.rerun()
                except Exception as e:
                    st.error(f"Error processing children CSV: {str(e)}")

        # Add new section for bulk parent upload
        with st.expander("Bulk Add Parents (CSV)"):
            st.info("Upload a CSV file with columns: 'name', 'email', 'password'")
            parent_upload = st.file_uploader("Upload Parents CSV", type=["csv"], key="parent_upload")

            if parent_upload:
                try:
                    df = pd.read_csv(parent_upload)

                    if not all(col in df.columns for col in ['name', 'email', 'password']):
                        st.error("CSV must contain 'name', 'email', and 'password' columns")
                    else:
                        st.write("Preview of parents to be added:")
                        st.dataframe(df.head())

                        if st.button("Add Parents"):
                            parents_data = []
                            for _, row in df.iterrows():
                                parent_data = {
                                    "id": str(uuid.uuid4()),
                                    "name": row['name'],
                                    "email": row['email'].strip().lower(),
                                    "password": row['password'],  # Note: In production, hash passwords
                                    "role": "Parent"
                                }
                                parents_data.append(parent_data)

                            # Insert in batches
                            batch_size = 50
                            for i in range(0, len(parents_data), batch_size):
                                batch = parents_data[i:i + batch_size]
                                supabase.table('users').insert(batch).execute()

                            st.success(f"Successfully added {len(parents_data)} parents!")
                            st.rerun()
                except Exception as e:
                    st.error(f"Error processing parents CSV: {str(e)}")

        # Add new section for bulk parent-child mapping upload
        with st.expander("Bulk Add Parent-Child Relationships (CSV)"):
            st.info("Upload a CSV file with columns: 'parent_email' and 'child_name'")
            mapping_upload = st.file_uploader("Upload Parent-Child CSV", type=["csv"], key="mapping_upload")

            if mapping_upload:
                try:
                    df = pd.read_csv(mapping_upload)

                    if not all(col in df.columns for col in ['parent_email', 'child_name']):
                        st.error("CSV must contain 'parent_email' and 'child_name' columns")
                    else:
                        st.write("Preview of relationships to be added:")
                        st.dataframe(df.head())

                        if st.button("Add Parent-Child Relationships"):
                            # Get all children and parents first
                            children = supabase.table('children').select("*").execute().data
                            parents = supabase.table('users').select("*").eq("role", "Parent").execute().data

                            child_name_to_id = {c['name'].lower(): c['id'] for c in children}
                            parent_email_to_id = {p['email'].lower(): p['id'] for p in parents}

                            success_count = 0
                            failed_rows = []

                            for idx, row in df.iterrows():
                                parent_email = row['parent_email'].strip().lower()
                                child_name = row['child_name'].strip().lower()

                                parent_id = parent_email_to_id.get(parent_email)
                                child_id = child_name_to_id.get(child_name)

                                if parent_id and child_id:
                                    # Update parent record with child_id
                                    supabase.table('users').update({'child_id': child_id}).eq('id', parent_id).execute()
                                    success_count += 1
                                else:
                                    failed_rows.append({
                                        "row": idx + 1,
                                        "parent_email": parent_email,
                                        "child_name": child_name,
                                        "error": ""
                                    })
                                    if not parent_id:
                                        failed_rows[-1]["error"] += "Parent not found. "
                                    if not child_id:
                                        failed_rows[-1]["error"] += "Child not found."

                            st.success(f"Successfully mapped {success_count} relationships!")

                            if failed_rows:
                                st.warning(f"Failed to process {len(failed_rows)} rows:")
                                st.dataframe(pd.DataFrame(failed_rows))
                            st.rerun()
                except Exception as e:
                    st.error(f"Error processing relationships CSV: {str(e)}")

    with tabs[2]:  # Observer-Child Mappings
        st.subheader("Observer-Child Mappings")

        # Add new section for CSV bulk upload
        with st.expander("Bulk Upload Observer-Student Mappings (CSV)"):
            st.info("Upload a CSV file with two columns: 'observer_id' and 'student_id'")
            uploaded_file = st.file_uploader("Choose CSV file", type=["csv"])

            if uploaded_file is not None:
                try:
                    # Read the CSV file
                    df = pd.read_csv(uploaded_file)

                    # Validate columns
                    if not all(col in df.columns for col in ['observer_id', 'student_id']):
                        st.error("CSV must contain 'observer_id' and 'student_id' columns")
                    else:
                        # Display preview
                        st.write("Preview of uploaded data:")
                        st.dataframe(df.head())

                        if st.button("Process CSV and Create Mappings"):
                            # Get all valid observers and students from database
                            observers = supabase.table('users').select("id").eq("role", "Observer").execute().data
                            valid_observer_ids = {o['id'] for o in observers}

                            children = supabase.table('children').select("id").execute().data
                            valid_child_ids = {c['id'] for c in children}

                            # Validate each row
                            valid_rows = []
                            invalid_rows = []

                            for idx, row in df.iterrows():
                                observer_id = row['observer_id']
                                child_id = row['student_id']

                                if observer_id in valid_observer_ids and child_id in valid_child_ids:
                                    valid_rows.append({
                                        "observer_id": observer_id,
                                        "child_id": child_id
                                    })
                                else:
                                    invalid_rows.append({
                                        "row": idx + 1,
                                        "observer_id": observer_id,
                                        "child_id": child_id,
                                        "error": ""
                                    })
                                    if observer_id not in valid_observer_ids:
                                        invalid_rows[-1]["error"] += "Invalid observer ID. "
                                    if child_id not in valid_child_ids:
                                        invalid_rows[-1]["error"] += "Invalid student ID."

                            # Insert valid mappings
                            if valid_rows:
                                result = supabase.table('observer_child_mappings').insert(valid_rows).execute()
                                st.success(f"Successfully added {len(valid_rows)} mappings!")

                            # Show invalid rows if any
                            if invalid_rows:
                                st.warning(f"{len(invalid_rows)} rows could not be processed:")
                                invalid_df = pd.DataFrame(invalid_rows)
                                st.dataframe(invalid_df)
                except Exception as e:
                    st.error(f"Error processing CSV: {str(e)}")

        try:
            mappings = supabase.table('observer_child_mappings').select("*").execute().data
            if mappings:
                for mapping in mappings:
                    # Get observer and child names
                    observer = supabase.table('users').select("*").eq('id', mapping['observer_id']).execute().data
                    child = supabase.table('children').select("*").eq('id', mapping['child_id']).execute().data

                    observer_name = observer[0].get('name', 'N/A') if observer else mapping['observer_id']
                    child_name = child[0].get('name', 'N/A') if child else mapping['child_id']

                    col1, col2, col3 = st.columns([4, 3, 1])
                    with col1:
                        st.write(f"Observer: {observer_name}")
                    with col2:
                        st.write(f"Child: {child_name}")
                    with col3:
                        if st.button("Delete", key=f"delete_{mapping['id']}"):
                            supabase.table('observer_child_mappings').delete().eq('id', mapping['id']).execute()
                            st.success("Mapping deleted successfully!")
                            st.rerun()
            else:
                st.info("No mappings found")

            with st.expander("Add New Observer-Child Mapping"):
                with st.form("add_observer_child"):
                    # Get all observers
                    observers = supabase.table('users').select("*").eq('role', 'Observer').execute().data
                    observer_options = {o['id']: f"{o.get('name', 'N/A')} ({o['email']})" for o in observers}

                    # Get all children
                    children = supabase.table('children').select("*").execute().data
                    child_options = {c['id']: c.get('name', 'N/A') for c in children}

                    observer_id = st.selectbox("Select Observer", options=list(observer_options.keys()),
                                               format_func=lambda x: observer_options[x])
                    child_id = st.selectbox("Select Child", options=list(child_options.keys()),
                                            format_func=lambda x: child_options[x])

                    if st.form_submit_button("Assign"):
                        if observer_id and child_id:
                            supabase.table('observer_child_mappings').insert({
                                "observer_id": observer_id,
                                "child_id": child_id
                            }).execute()
                            st.success("Mapping created successfully!")
                            st.rerun()
        except Exception as e:
            st.error(f"Database error: {str(e)}")

    with tabs[3]:  # Activity Logs
        st.subheader("Activity Logs")
        try:
            logs = supabase.table('observer_activity_logs').select("*").order('timestamp', desc=True).limit(
                100).execute().data
            if logs:
                st.dataframe(logs)
            else:
                st.info("No activity logs")
        except Exception as e:
            st.error(f"Database error: {str(e)}")

    with tabs[4]:  # Report Processing tab
        st.subheader("Process Reports for Observers")

        # Initialize session state variables for admin report processing
        if 'admin_processing_mode' not in st.session_state:
            st.session_state.admin_processing_mode = None
        if 'admin_report_generated' not in st.session_state:
            st.session_state.admin_report_generated = None
        if 'admin_transcript' not in st.session_state:
            st.session_state.admin_transcript = ""

        # Get all observers
        observers = supabase.table('users').select("*").eq("role", "Observer").execute().data
        if not observers:
            st.warning("No observers found in the system")
        else:
            # Select an observer
            observer_options = {o['id']: f"{o.get('name', 'N/A')} ({o['email']})" for o in observers}
            selected_observer_id = st.selectbox(
                "Select Observer",
                options=list(observer_options.keys()),
                format_func=lambda x: observer_options[x],
                key="admin_observer_select"
            )

            # Get children assigned to selected observer
            mappings = supabase.table('observer_child_mappings').select("child_id").eq("observer_id",
                                                                                       selected_observer_id).execute().data
            child_ids = [m['child_id'] for m in mappings]

            if not child_ids:
                st.warning("No children assigned to this observer")
            else:
                # Get child details
                children = supabase.table('children').select("*").in_("id", child_ids).execute().data
                child_options = {c['id']: c.get('name', f"Child {c['id'][:4]}") for c in children}

                # Select a child
                selected_child_id = st.selectbox(
                    "Select Student",
                    options=list(child_options.keys()),
                    format_func=lambda x: child_options[x],
                    key="admin_child_select"
                )

                # Initialize the extractor
                extractor = ObservationExtractor()

                # Choose processing mode
                st.subheader("Select Processing Mode")
                col1, col2 = st.columns(2)

                with col1:
                    if st.button("OCR Mode (Image Upload)", key="admin_ocr_btn"):
                        st.session_state.admin_processing_mode = "ocr"
                        st.session_state.admin_report_generated = None
                        st.session_state.admin_transcript = ""
                        st.rerun()
                with col2:
                    if st.button("Audio Mode (Recording Upload)", key="admin_audio_btn"):
                        st.session_state.admin_processing_mode = "audio"
                        st.session_state.admin_report_generated = None
                        st.session_state.admin_transcript = ""
                        st.rerun()

                # Session information
                st.subheader("Session Information")
                observer_name = observer_options[selected_observer_id].split(" (")[0]
                student_name = child_options[selected_child_id]

                session_date = st.date_input("Session Date:", key="admin_session_date").strftime('%d/%m/%Y')
                col1, col2 = st.columns(2)
                with col1:
                    session_start = st.time_input("Start Time:", key="admin_start_time").strftime('%H:%M')
                with col2:
                    session_end = st.time_input("End Time:", key="admin_end_time").strftime('%H:%M')

                user_info = {
                    'student_name': student_name,
                    'observer_name': observer_name,
                    'session_date': session_date,
                    'session_start': session_start,
                    'session_end': session_end
                }

                # OCR Processing
                if st.session_state.admin_processing_mode == "ocr":
                    st.info("OCR Mode: Upload an image of an observation sheet")
                    uploaded_file = st.file_uploader("Upload Observation Sheet", type=["jpg", "jpeg", "png"],
                                                     key="admin_ocr_upload")

                    if uploaded_file and st.button("Process Observation", key="admin_process_ocr"):
                        with st.spinner("Processing..."):
                            try:
                                extracted_text = extractor.extract_text_with_ocr(uploaded_file)
                                structured_data = extractor.process_with_groq(extracted_text)
                                observations_text = structured_data.get("observations", "")

                                # Upload file to storage
                                file_url = None
                                if uploaded_file:
                                    # Reset file pointer to beginning
                                    uploaded_file.seek(0)
                                    # Upload to storage
                                    file_url = upload_file_to_storage(
                                        uploaded_file.getvalue(),
                                        uploaded_file.name,
                                        f"image/{uploaded_file.type.split('/')[1]}"
                                    )

                                if observations_text:
                                    # Generate report with new format
                                    report = extractor.generate_report_from_text(observations_text, user_info)
                                    st.session_state.admin_report_generated = report

                                    # Insert observation
                                    observation_response = supabase.table('observations').insert({
                                        "student_id": selected_child_id,
                                        "username": selected_observer_id,  # Using the selected observer's ID
                                        "student_name": structured_data.get("studentName", student_name),
                                        "observer_name": observer_name,
                                        "class_name": structured_data.get("className", ""),
                                        "date": structured_data.get("date", session_date),
                                        "observations": observations_text,
                                        "strengths": json.dumps(structured_data.get("strengths", [])),
                                        "areas_of_development": json.dumps(
                                            structured_data.get("areasOfDevelopment", [])),
                                        "recommendations": json.dumps(structured_data.get("recommendations", [])),
                                        "timestamp": datetime.now().isoformat(),
                                        "filename": uploaded_file.name,
                                        "full_data": json.dumps(structured_data),
                                        "theme_of_day": structured_data.get("themeOfDay", ""),
                                        "curiosity_seed": structured_data.get("curiositySeed", ""),
                                        "processed_by_admin": True,  # Flag to indicate admin processed this
                                        "file_url": file_url  # Add the file URL
                                    }).execute()

                                    st.success("Data processed and saved successfully!")
                                else:
                                    st.error("No observations found in the extracted data")
                            except Exception as e:
                                st.error(f"Processing error: {str(e)}")


                # Audio Processing
                elif st.session_state.admin_processing_mode == "audio":
                    st.info("Audio Mode: Upload an audio recording of an observation session")
                    uploaded_file = st.file_uploader(
                        "Choose an audio file",
                        type=["wav", "mp3", "m4a", "mpeg", "ogg", "flac", "aac", "wma", "aiff"],
                        key="admin_audio_upload"
                    )

                    if uploaded_file and st.button("Process & Generate Report", key="admin_process_audio"):
                        if not assemblyai_key:
                            st.error("AssemblyAI API key is missing.")
                        else:
                            with st.spinner("Step 1/2: Transcribing audio..."):
                                transcript = extractor.transcribe_with_assemblyai(uploaded_file)
                                st.session_state.admin_transcript = transcript

                            with st.spinner("Step 2/2: Generating report..."):
                                # Upload file to storage
                                file_url = None
                                if uploaded_file:
                                    # Reset file pointer to beginning
                                    uploaded_file.seek(0)
                                    # Upload to storage
                                    file_url = upload_file_to_storage(
                                        uploaded_file.getvalue(),
                                        uploaded_file.name,
                                        f"audio/{uploaded_file.type.split('/')[1]}"
                                    )

                                # Generate report with new format
                                report = extractor.generate_report_from_text(transcript, user_info)
                                st.session_state.admin_report_generated = report

                                # Extract theme and curiosity seed using AI
                                try:
                                    theme_prompt = f"""
                                    Based on this transcript, identify:
                                    1. The main theme or topic of the day
                                    2. A curiosity seed (something that sparked the child's interest or curiosity)

                                    Return as JSON with keys: "themeOfDay" and "curiositySeed"

                                    TRANSCRIPT:
                                    {transcript}
                                    """

                                    model = genai.GenerativeModel('gemini-2.0-flash-002')
                                    response = model.generate_content([
                                        {"role": "user", "parts": [{"text": theme_prompt}]}
                                    ])

                                    theme_data = json.loads(response.text)
                                    theme_of_day = theme_data.get("themeOfDay", "")
                                    curiosity_seed = theme_data.get("curiositySeed", "")
                                except:
                                    theme_of_day = ""
                                    curiosity_seed = ""

                                # Save to database
                                supabase.table('observations').insert({
                                    "student_id": selected_child_id,
                                    "username": selected_observer_id,  # Using the selected observer's ID
                                    "student_name": student_name,
                                    "observer_name": observer_name,
                                    "class_name": "",
                                    "date": session_date,
                                    "observations": transcript,
                                    "strengths": json.dumps([]),
                                    "areas_of_development": json.dumps([]),
                                    "recommendations": json.dumps([]),
                                    "timestamp": datetime.now().isoformat(),
                                    "filename": uploaded_file.name,
                                    "full_data": json.dumps({"transcript": transcript, "report": report}),
                                    "theme_of_day": theme_of_day,
                                    "curiosity_seed": curiosity_seed,
                                    "processed_by_admin": True,  # Flag to indicate admin processed this
                                    "file_url": file_url  # Add the file URL
                                }).execute()

                                st.success("Audio processed and report generated successfully!")

                # Report Display and Download
                if st.session_state.admin_report_generated:
                    st.subheader("Generated Report")
                    st.markdown(st.session_state.admin_report_generated)
                    docx_file = extractor.create_word_document(st.session_state.admin_report_generated)
                    filename = f"Observer_Report_{student_name.replace(' ', '_')}_{session_date.replace('/', '-')}.docx"
                    st.download_button(
                        "Download as Word Document",
                        data=docx_file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    # Email option
                    st.subheader("Email Report")
                    with st.form("admin_email_form"):
                        to_email = st.text_input("Recipient Email", value="parent@example.com")
                        subject = st.text_input("Subject", value=f"Observer Report for {student_name}")
                        submitted = st.form_submit_button("Send Email")
                        if submitted:
                            success, message = extractor.send_email(to_email, subject,
                                                                    st.session_state.admin_report_generated)
                            if success:
                                st.success(message)
                            else:
                                st.error(message)

                    # Show transcript for audio processing
                    if st.session_state.admin_transcript:
                        with st.expander("View Transcript"):
                            st.text_area("Transcript", value=st.session_state.admin_transcript, height=300,
                                         disabled=True)


# Parent Dashboard
def parent_dashboard(user_id):
    st.title(f"Parent Portal")

    # Create tabs for different sections
    parent_tabs = st.tabs(["Reports", "Messages", "Goals"])

    with parent_tabs[0]:  # Reports tab
        try:
            # Get the parent's information
            parent_data = supabase.table('users').select("*").eq("id", user_id).execute().data
            if not parent_data:
                st.warning("User not found")
                return

            parent = parent_data[0]
            child_id = parent.get('child_id')

            if not child_id:
                st.warning("No child assigned to your account. Please contact admin.")
                return

            # Get child information
            child_data = supabase.table('children').select("*").eq("id", child_id).execute().data
            if not child_data:
                st.warning("Child information not found")
                return

            child = child_data[0]

            # Get observer information
            mapping = supabase.table('observer_child_mappings').select("*").eq("child_id", child_id).execute().data
            observer_id = mapping[0]['observer_id'] if mapping else None
            observer_name = "Not assigned"

            if observer_id:
                observer_data = supabase.table('users').select("name").eq("id", observer_id).execute().data
                if observer_data:
                    observer_name = observer_data[0].get('name', observer_id)

            # Display dashboard
            st.subheader(f"Your Child: {child.get('name', 'N/A')}")

            cols = st.columns(3)
            cols[0].metric("Age", child.get('age', 'N/A'))
            cols[1].metric("Grade/Class", child.get('grade', 'N/A'))
            cols[2].metric("Assigned Observer", observer_name)

            st.subheader("Recent Reports")
            reports = supabase.table('observations').select("*").eq("student_id", child_id).order('date',
                                                                                                  desc=True).execute().data

            if reports:
                for report in reports:
                    with st.expander(f"Report from {report.get('date', 'unknown date')}"):
                        st.write(f"**Observer:** {report.get('observer_name', 'N/A')}")
                        st.write(f"**Date:** {report.get('date', 'N/A')}")

                        if report.get('observations'):
                            st.write("**Observations:**")
                            st.write(report['observations'])

                        if report.get('strengths'):
                            try:
                                strengths = json.loads(report['strengths']) if isinstance(report['strengths'], str) else \
                                    report['strengths']
                                st.write("**Strengths:**")
                                for strength in strengths:
                                    st.write(f"- {strength}")
                            except:
                                pass

                        if report.get('areas_of_development'):
                            try:
                                areas = json.loads(report['areas_of_development']) if isinstance(
                                    report['areas_of_development'], str) else report['areas_of_development']
                                st.write("**Areas for Development:**")
                                for area in areas:
                                    st.write(f"- {area}")
                            except:
                                pass

                        if report.get('recommendations'):
                            try:
                                recs = json.loads(report['recommendations']) if isinstance(report['recommendations'],
                                                                                           str) else report[
                                    'recommendations']
                                st.write("**Recommendations:**")
                                for rec in recs:
                                    st.write(f"- {rec}")
                            except:
                                pass
            else:
                st.info("No reports available yet")

            st.markdown("---")
            monthly_report_section(child_id, user_id)

        except Exception as e:
            st.error(f"Database error: {str(e)}")

    with parent_tabs[1]:  # Messages tab
        st.subheader("Messages")

        # Get child and observer info
        parent_data = supabase.table('users').select("*").eq("id", user_id).execute().data
        if not parent_data:
            st.warning("User not found")
            return

        parent = parent_data[0]
        child_id = parent.get('child_id')

        if not child_id:
            st.warning("No child assigned to your account. Please contact admin.")
            return

        # Get observer information
        mapping = supabase.table('observer_child_mappings').select("*").eq("child_id", child_id).execute().data
        observer_id = mapping[0]['observer_id'] if mapping else None

        if not observer_id:
            st.warning("No observer assigned to your child yet.")
            return

        # Get observer details
        observer_data = supabase.table('users').select("*").eq("id", observer_id).execute().data
        if not observer_data:
            st.warning("Observer information not found")
            return

        observer = observer_data[0]

        # Display messaging interface
        st.write(f"**Messaging with:** {observer.get('name', 'Observer')}")

        # Get existing messages
        messages = supabase.table('messages').select("*") \
            .or_(f"sender_id.eq.{user_id},receiver_id.eq.{user_id}") \
            .or_(f"sender_id.eq.{observer_id},receiver_id.eq.{observer_id}") \
            .order('timestamp', desc=False) \
            .execute().data

        # Display message history
        message_container = st.container(height=400)
        with message_container:
            for msg in messages:
                is_from_me = msg['sender_id'] == user_id

                col1, col2 = st.columns([1, 4])
                with col1:
                    st.write("You:" if is_from_me else f"{observer.get('name', 'Observer')}:")
                with col2:
                    st.text_area("", value=msg['content'], height=70, key=f"msg_{msg['id']}", disabled=True)

        # Send new message
        with st.form("send_message_form"):
            new_message = st.text_area("Type your message:", height=100)
            if st.form_submit_button("Send Message"):
                if new_message.strip():
                    # Insert message into database
                    supabase.table('messages').insert({
                        "sender_id": user_id,
                        "receiver_id": observer_id,
                        "content": new_message,
                        "timestamp": datetime.now().isoformat(),
                        "read": False
                    }).execute()
                    st.success("Message sent!")
                    st.rerun()
                else:
                    st.error("Cannot send empty message")

    with parent_tabs[2]:  # Goals tab
        st.markdown("---")
        st.subheader("Goal Tracking")

        parent_data = supabase.table('users').select("*").eq("id", user_id).execute().data
        if not parent_data:
            st.warning("User not found")
            return

        parent = parent_data[0]
        child_id = parent.get('child_id')

        if child_id:
            goals = supabase.table('goals').select("*").eq("child_id", child_id).execute().data

            if goals:
                for goal in goals:
                    with st.expander(f"Goal from {goal.get('created_at', 'unknown date')}"):
                        st.write(goal['goal_text'])
                        st.write(f"Status: {goal.get('status', 'active')}")
                        st.write(f"Target Date: {goal.get('target_date', 'No target date')}")

                        # Show alignments with reports
                        alignments = supabase.table('goal_alignments').select("*").eq("goal_id",
                                                                                      goal['id']).execute().data
                        if alignments:
                            st.write("**Report Alignments:**")
                            for alignment in alignments:
                                report = supabase.table('observations').select("date").eq("id", alignment[
                                    'report_id']).execute().data
                                report_date = report[0]['date'] if report else "Unknown date"
                                st.write(f"- {report_date}: Score {alignment.get('alignment_score', 0)}/10")

                                # Feedback form if no feedback exists for this alignment
                                existing_feedback = supabase.table('parent_feedback').select("*").eq("alignment_id",
                                                                                                     alignment[
                                                                                                         'id']).eq(
                                    "parent_id", user_id).execute().data

                                if not existing_feedback:
                                    with st.form(f"feedback_form_{alignment['id']}"):
                                        rating = st.slider("Rate this alignment", 1, 5, 3,
                                                           key=f"rating_{alignment['id']}")
                                        feedback_text = st.text_area("Your feedback", key=f"feedback_{alignment['id']}")

                                        if st.form_submit_button("Submit Feedback"):
                                            feedback_data = {
                                                "alignment_id": alignment['id'],
                                                "parent_id": user_id,
                                                "feedback_text": feedback_text,
                                                "rating": rating
                                            }
                                            supabase.table('parent_feedback').insert(feedback_data).execute()
                                            st.success("Feedback submitted!")
                                            st.rerun()
                                else:
                                    fb = existing_feedback[0]
                                    st.write(f"**Your Feedback:** {'⭐' * fb.get('rating', 0)}")
                                    st.write(fb.get('feedback_text', 'No feedback text'))
            else:
                st.info("No goals set for your child yet")


def monthly_report_section(child_id, parent_id):
    """Display monthly report section for parents"""
    st.subheader("Monthly Progress Reports")

    # Initialize the report generator
    report_generator = MonthlyReportGenerator(supabase)

    # Get child's info
    child_data = supabase.table('children').select("*").eq("id", child_id).execute().data
    if not child_data:
        st.warning("Child information not found")
        return

    child = child_data[0]

    # Date selection
    col1, col2 = st.columns(2)
    with col1:
        current_date = datetime.now()
        year = st.selectbox("Year",
                            options=list(range(current_date.year - 2, current_date.year + 1)),
                            index=2)  # Default to current year
    with col2:
        month = st.selectbox("Month",
                             options=list(range(1, 13)),
                             format_func=lambda x: calendar.month_name[x],
                             index=current_date.month - 1)  # Default to current month

    # Check if report already exists
    existing_report = supabase.table('monthly_reports').select("*") \
        .eq("child_id", child_id) \
        .eq("parent_id", parent_id) \
        .eq("month", month) \
        .eq("year", year) \
        .execute().data

    # Get all observations for this child in the selected month
    observations = report_generator.get_month_data(child_id, year, month)

    if not observations:
        st.info(f"No observations found for {calendar.month_name[month]} {year}")
        return

    # Get goal progress for this month
    goal_progress = report_generator.get_goal_progress(child_id, year, month)

    # Display monthly summary
    summary = report_generator.generate_monthly_summary(observations, goal_progress)
    st.markdown(summary)

    # Display observation frequency chart
    obs_freq_chart = report_generator.generate_observation_frequency_chart(observations)
    if obs_freq_chart:
        st.plotly_chart(obs_freq_chart, use_container_width=True)

    # Extract strengths and development areas
    strength_counts = report_generator.get_strength_areas(observations)
    development_counts = report_generator.get_development_areas(observations)

    # Display strengths and development areas
    col1, col2 = st.columns(2)

    with col1:
        strengths_chart = report_generator.generate_strengths_chart(strength_counts)
        if strengths_chart:
            st.plotly_chart(strengths_chart, use_container_width=True)
        else:
            st.info("No strengths data available")

    with col2:
        development_chart = report_generator.generate_development_areas_chart(development_counts)
        if development_chart:
            st.plotly_chart(development_chart, use_container_width=True)
        else:
            st.info("No development areas data available")

    # Display goal progress
    if goal_progress:
        goal_chart = report_generator.generate_goal_progress_chart(goal_progress)
        if goal_chart:
            st.plotly_chart(goal_chart, use_container_width=True)
    else:
        st.info("No goal progress data available for this month")

    # Prepare report data
    report_data = {
        "summary": summary,
        "strength_counts": strength_counts,
        "development_counts": development_counts,
        "goal_progress": goal_progress,
        "num_observations": len(observations)
    }

    # Save report if it doesn't exist
    if not existing_report:
        if st.button("Save Monthly Report"):
            report_id = str(uuid.uuid4())
            supabase.table('monthly_reports').insert({
                "id": report_id,
                "child_id": child_id,
                "parent_id": parent_id,
                "month": month,
                "year": year,
                "report_data": report_data
            }).execute()
            st.success("Monthly report saved successfully!")
            st.rerun()
    else:
        report = existing_report[0]
        st.success("Monthly report already generated for this period")

        # Show feedback section if no feedback exists
        if not report.get('feedback'):
            with st.form("monthly_feedback"):
                st.subheader("Provide Feedback")
                feedback = st.text_area("Your feedback on this monthly report")
                rating = st.slider("Rating (1-5)", 1, 5, 3)

                if st.form_submit_button("Submit Feedback"):
                    supabase.table('monthly_reports').update({
                        "feedback": feedback,
                        "feedback_submitted_at": datetime.now().isoformat()
                    }).eq("id", report['id']).execute()
                    st.success("Thank you for your feedback!")
                    st.rerun()
        else:
            st.subheader("Your Feedback")
            st.write(f"Rating: {'⭐' * report.get('rating', 0)}")
            st.write(report['feedback'])

    # Add download option for report
    if st.button("Generate Downloadable Report"):
        # Create DataFrame with main metrics
        report_data = {
            "Metric": ["Total Observations", "Goals Tracked", "Average Goal Score"],
            "Value": [len(observations), len(goal_progress),
                      sum(g['avg_score'] for g in goal_progress) / len(goal_progress) if goal_progress else 0]
        }

        df = pd.DataFrame(report_data)

        # Create Excel buffer
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name='Summary', index=False)

            # Add strengths sheet
            if strength_counts:
                strengths_df = pd.DataFrame([
                    {"Strength": strength, "Count": count}
                    for strength, count in strength_counts.items()
                ])
                strengths_df.to_excel(writer, sheet_name='Strengths', index=False)

            # Add development areas sheet
            if development_counts:
                development_df = pd.DataFrame([
                    {"Development Area": area, "Count": count}
                    for area, count in development_counts.items()
                ])
                development_df.to_excel(writer, sheet_name='Development Areas', index=False)

            # Add goal progress sheet
            if goal_progress:
                goals_df = pd.DataFrame([
                    {"Goal": g['goal_text'], "Average Score": g['avg_score'], "Observations": g['num_observations']}
                    for g in goal_progress
                ])
                goals_df.to_excel(writer, sheet_name='Goal Progress', index=False)

        buffer.seek(0)

        # Create download button
        month_name = calendar.month_name[month]
        st.download_button(
            label="Download Excel Report",
            data=buffer,
            file_name=f"{child['name']}_Progress_Report_{month_name}_{year}.xlsx",
            mime="application/vnd.ms-excel"
        )


def observer_monthly_report_section(observer_id):
    """Display monthly report section for observers"""
    st.subheader("Monthly Progress Reports")

    # Initialize the report generator
    report_generator = MonthlyReportGenerator(supabase)

    # Get all children assigned to this observer
    mappings = supabase.table('observer_child_mappings').select("child_id").eq("observer_id",
                                                                               observer_id).execute().data
    child_ids = [m['child_id'] for m in mappings]

    if not child_ids:
        st.warning("No children assigned to you yet")
        return

    # Get child details
    children = supabase.table('children').select("*").in_("id", child_ids).execute().data
    child_options = {c['id']: c.get('name', f"Child {c['id'][:4]}") for c in children}

    # Select child
    selected_child_id = st.selectbox(
        "Select Student",
        options=list(child_options.keys()),
        format_func=lambda x: child_options[x],
        key="audio_child_select"
    )

    # Date selection
    col1, col2 = st.columns(2)
    with col1:
        # Get the current year and month
        current_date = datetime.now()
        year = st.selectbox("Year",
                            options=list(range(current_date.year - 2, current_date.year + 1)),
                            index=2)  # Default to current year
    with col2:
        month = st.selectbox("Month",
                             options=list(range(1, 13)),
                             format_func=lambda x: calendar.month_name[x],
                             index=current_date.month - 1)  # Default to current month

    # Get all observations for this child in the selected month
    observations = report_generator.get_month_data(selected_child_id, year, month)

    if not observations:
        st.info(f"No observations found for {calendar.month_name[month]} {year}")
        return

    # Get goal progress for this month
    goal_progress = report_generator.get_goal_progress(selected_child_id, year, month)

    # Display monthly summary
    summary = report_generator.generate_monthly_summary(observations, goal_progress)
    st.markdown(summary)

    # Display observation frequency chart
    obs_freq_chart = report_generator.generate_observation_frequency_chart(observations)
    if obs_freq_chart:
        st.plotly_chart(obs_freq_chart, use_container_width=True)

    # Extract strengths and development areas
    strength_counts = report_generator.get_strength_areas(observations)
    development_counts = report_generator.get_development_areas(observations)

    # Display strengths and development areas
    col1, col2 = st.columns(2)

    with col1:
        strengths_chart = report_generator.generate_strengths_chart(strength_counts)
        if strengths_chart:
            st.plotly_chart(strengths_chart, use_container_width=True)
        else:
            st.info("No strengths data available")

    with col2:
        development_chart = report_generator.generate_development_areas_chart(development_counts)
        if development_chart:
            st.plotly_chart(development_chart, use_container_width=True)
        else:
            st.info("No development areas data available")

    # Display goal progress
    if goal_progress:
        goal_chart = report_generator.generate_goal_progress_chart(goal_progress)
        if goal_chart:
            st.plotly_chart(goal_chart, use_container_width=True)
    else:
        st.info("No goal progress data available for this month")

    # Add option to share report with parents
    if st.button("Share Report with Parent"):
        try:
            # Find parent associated with this child
            parent = supabase.table('users').select("*").eq("child_id", selected_child_id).eq("role",
                                                                                              "Parent").execute().data

            if parent:
                parent = parent[0]
                parent_email = parent['email']

                # Create email content
                email_subject = f"Monthly Progress Report - {child_options[selected_child_id]} - {calendar.month_name[month]} {year}"
                email_body = f"""
                <h2>Monthly Progress Report</h2>
                <p><strong>Child:</strong> {child_options[selected_child_id]}</p>
                <p><strong>Period:</strong> {calendar.month_name[month]} {year}</p>
                <p><strong>Observations:</strong> {len(observations)}</p>
                <p><strong>Goals Tracked:</strong> {len(goal_progress)}</p>

                <p>Please log in to the Learning Observer platform to view the full report with charts and details.</p>
                """

                # Send email
                extractor = ObservationExtractor()  # Create an instance of the class to use the send_email method
                success, message = extractor.send_email(parent_email, email_subject, email_body)

                if success:
                    st.success(f"Report shared successfully with {parent_email}")
                else:
                    st.error(f"Failed to share report: {message}")
            else:
                st.warning("No parent found for this child")
        except Exception as e:
            st.error(f"Error sharing report: {str(e)}")

    # Add download option for report
    if st.button("Download Report"):
        # Create DataFrame with main metrics
        report_data = {
            "Metric": ["Total Observations", "Goals Tracked", "Average Goal Score"],
            "Value": [len(observations), len(goal_progress),
                      sum(g['avg_score'] for g in goal_progress) / len(goal_progress) if goal_progress else 0]
        }

        df = pd.DataFrame(report_data)

        # Create Excel buffer
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name='Summary', index=False)

            # Add strengths sheet
            if strength_counts:
                strengths_df = pd.DataFrame([
                    {"Strength": strength, "Count": count}
                    for strength, count in strength_counts.items()
                ])
                strengths_df.to_excel(writer, sheet_name='Strengths', index=False)

            # Add development areas sheet
            if development_counts:
                development_df = pd.DataFrame([
                    {"Development Area": area, "Count": count}
                    for area, count in development_counts.items()
                ])
                development_df.to_excel(writer, sheet_name='Development Areas', index=False)

            # Add goal progress sheet
            if goal_progress:
                goals_df = pd.DataFrame([
                    {"Goal": g['goal_text'], "Average Score": g['avg_score'], "Observations": g['num_observations']}
                    for g in goal_progress
                ])
                goals_df.to_excel(writer, sheet_name='Goal Progress', index=False)

        buffer.seek(0)

        # Create download button
        month_name = calendar.month_name[month]
        st.download_button(
            label="Download Excel Report",
            data=buffer,
            file_name=f"{child_options[selected_child_id]}_Progress_Report_{month_name}_{year}.xlsx",
            mime="application/vnd.ms-excel"
        )


# Main App
def main():
    extractor = ObservationExtractor()

    # Define AssemblyAI API key
    assemblyai_key = st.secrets.get("ASSEMBLYAI_API_KEY", "")

    # Session State Initialization
    if 'auth' not in st.session_state:
        st.session_state.auth = {
            'logged_in': False,
            'role': None,
            'user_id': None,
            'email': None,
            'name': None
        }
    if 'user_info' not in st.session_state:
        st.session_state.user_info = {
            'student_name': '',
            'observer_name': '',
            'session_date': datetime.now().strftime('%d/%m/%Y'),
            'session_start': '',
            'session_end': ''
        }
    if 'audio_transcription' not in st.session_state:
        st.session_state.audio_transcription = ""
    if 'report_generated' not in st.session_state:
        st.session_state.report_generated = None
    if 'show_edit_transcript' not in st.session_state:
        st.session_state.show_edit_transcript = False
    if 'processing_mode' not in st.session_state:
        st.session_state.processing_mode = None
    if 'show_register' not in st.session_state:
        st.session_state.show_register = False
    if 'admin_initial_login' not in st.session_state:
        st.session_state.admin_initial_login = True

    # Admin credentials
    ADMIN_CREDS = {
        "username": st.secrets.get("ADMIN_USER", "admin"),
        "password": st.secrets.get("ADMIN_PASS", "hello")
    }

    # Login/Registration Page
    if not st.session_state.auth['logged_in']:
        st.title("Learning Observer Login")

        if st.session_state.show_register:
            # Registration Form with improved validation
            with st.form("register_form"):
                st.subheader("Create New Account")
                name = st.text_input("Full Name", max_chars=100)
                email = st.text_input("Email", max_chars=100).strip().lower()
                role = st.selectbox("Role", ["Observer", "Parent"])
                password = st.text_input("Password", type="password", max_chars=100)
                confirm_password = st.text_input("Confirm Password", type="password", max_chars=100)

                if role == "Parent":
                    # For parents, show a dropdown of available children
                    children = supabase.table('children').select("*").execute().data
                    child_options = {c['id']: c.get('id', 'N/A') for c in children}
                    child_id = st.selectbox("Select Your Child", options=list(child_options.keys()),
                                            format_func=lambda x: child_options[x])

                submitted = st.form_submit_button("Register")

                if submitted:
                    if not all([name, email, password, confirm_password]):
                        st.error("Please fill in all required fields")
                    elif password != confirm_password:
                        st.error("Passwords do not match!")
                    elif len(password) < 8:
                        st.error("Password must be at least 8 characters")
                    else:
                        try:
                            # Check if email exists
                            existing_response = supabase.table('users').select("email").eq("email", email).execute()

                            if existing_response.data:
                                st.error("Email already registered. Please login instead.")
                            else:
                                # Prepare user data with explicit column names
                                user_data = {
                                    "id": str(uuid.uuid4()),
                                    "email": email,
                                    "name": name,
                                    "password": password,  # Note: In production, hash this password!
                                    "role": role
                                }

                                if role == "Parent":
                                    user_data["child_id"] = child_id

                                # Insert with error handling
                                insert_response = supabase.table('users').insert(user_data).execute()

                                if insert_response.data:
                                    st.success("Registration successful! Please login.")
                                    st.session_state.show_register = False
                                else:
                                    st.error("Registration failed. Please try again.")

                        except Exception as e:
                            logger.error(f"Registration error: {str(e)}")
                            st.error(f"Registration failed: {str(e)}")

            if st.button("Back to Login"):
                st.session_state.show_register = False
        else:
            # Login Form with improved error handling
            with st.form("login_form"):
                email = st.text_input("Email", max_chars=100).strip().lower()
                password = st.text_input("Password", type="password", max_chars=100)

                submitted = st.form_submit_button("Login")

                if submitted:
                    try:
                        # Check admin login first
                        if email == ADMIN_CREDS["username"] and password == ADMIN_CREDS["password"]:
                            st.session_state.auth = {
                                'logged_in': True,
                                'role': 'Admin',
                                'user_id': 'admin',
                                'email': email,
                                'name': 'Admin'
                            }
                            st.rerun()

                        # Check regular user login
                        user_response = supabase.table('users').select("*").eq("email", email).eq("password",
                                                                                                  password).execute()

                        if user_response.data:
                            user = user_response.data[0]
                            st.session_state.auth = {
                                'logged_in': True,
                                'role': user['role'],
                                'user_id': user['id'],
                                'email': user['email'],
                                'name': user.get('name', 'User')
                            }
                            st.rerun()
                        else:
                            st.error("Invalid email or password")

                    except Exception as e:
                        logger.error(f"Login error: {str(e)}")
                        st.error(f"Login failed: {str(e)}")

            st.write("Don't have an account?")
            if st.button("Register Here"):
                st.session_state.show_register = True

        return

    # Logout Button (common)
    def logout_button():
        if st.button("Logout"):
            if st.session_state.auth['role'] == "Observer":
                supabase.table('observer_activity_logs').insert({
                    "observer_id": st.session_state.auth['user_id'],
                    "child_id": "N/A",
                    "action": "logout",
                    "duration_minutes": 0
                }).execute()
            st.session_state.auth = {'logged_in': False, 'role': None, 'user_id': None}
            st.rerun()

    # Admin Dashboard - Modified to check for initial login
    if st.session_state.auth['role'] == 'Admin':
        if st.session_state.admin_initial_login:
            # Show a welcome or intermediate page instead of the full admin dashboard
            st.title("Welcome, Admin")
            st.write("You are logged in as an administrator.")
            # Add a button to proceed to the dashboard
            if st.button("Access Admin Dashboard"):
                st.session_state.admin_initial_login = False
                st.rerun()

            logout_button()
            return
        else:
            # Show the regular admin dashboard
            admin_dashboard()
            logout_button()
            return

    # Parent Dashboard
    if st.session_state.auth['role'] == 'Parent':
        parent_dashboard(st.session_state.auth['user_id'])
        logout_button()
        return

    # Observer Dashboard
    st.title(f"Observer Dashboard - {st.session_state.auth['name']}")

    # Log login activity
    supabase.table('observer_activity_logs').insert({
        "observer_id": st.session_state.auth['user_id'],
        "child_id": "N/A",
        "action": "login",
        "duration_minutes": 0
    }).execute()

    logout_button()
    observer_tabs = st.tabs(["Observation Processing", "Goal Management", "Messages", "Monthly Reports"])

    with observer_tabs[0]:
        # Sidebar for user information
        with st.sidebar:
            st.subheader("Session Information")
            st.session_state.user_info['student_name'] = st.text_input("Student Name:",
                                                                       value=st.session_state.user_info['student_name'])
            st.session_state.user_info['observer_name'] = st.text_input("Observer Name:",
                                                                        value=st.session_state.user_info[
                                                                                  'observer_name'] or st.session_state.auth.get(
                                                                            'name', ''))
            st.session_state.user_info['session_date'] = st.date_input("Session Date:").strftime('%d/%m/%Y')
            col1, col2 = st.columns(2)
            with col1:
                st.session_state.user_info['session_start'] = st.time_input("Start Time:").strftime('%H:%M')
            with col2:
                st.session_state.user_info['session_end'] = st.time_input("End Time:").strftime('%H:%M')

        # Choose processing mode
        st.subheader("Select Processing Mode")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("OCR Mode (Image Upload)"):
                st.session_state.processing_mode = "ocr"
                st.session_state.audio_transcription = ""
                st.session_state.report_generated = None
        with col2:
            if st.button("Audio Mode (Recording Upload)"):
                st.session_state.processing_mode = "audio"
                st.session_state.audio_transcription = ""
                st.session_state.report_generated = None

        # Get children assigned to this observer
        mappings = supabase.table('observer_child_mappings').select("child_id").eq("observer_id", st.session_state.auth[
            'user_id']).execute().data
        child_ids = [m['child_id'] for m in mappings]
        children = supabase.table('children').select("*").in_("id", child_ids).execute().data
        child_options = {c['id']: c.get('name', f"Child {c['id'][:4]}") for c in children}

        selected_child_id = st.selectbox(
            "Select Student",
            options=list(child_options.keys()),
            format_func=lambda x: child_options[x],
            key="ocr_child_select"
        )

        # OCR Processing
        # OCR Processing
        if st.session_state.processing_mode == "ocr":
            st.info("OCR Mode: Upload an image of an observation sheet")
            uploaded_file = st.file_uploader("Upload Observation Sheet", type=["jpg", "jpeg", "png"])
            if uploaded_file and st.button("Process Observation"):
                with st.spinner("Processing..."):
                    try:
                        # Process the image with OCR
                        extracted_text = extractor.extract_text_with_ocr(uploaded_file)
                        structured_data = extractor.process_with_groq(extracted_text)
                        observations_text = structured_data.get("observations", "")

                        # Upload file to storage
                        file_url = None
                        if uploaded_file:
                            # Reset file pointer to beginning
                            uploaded_file.seek(0)
                            # Upload to storage
                            file_url = upload_file_to_storage(
                                uploaded_file.getvalue(),
                                uploaded_file.name,
                                f"image/{uploaded_file.type.split('/')[1]}"
                            )

                        if observations_text:
                            # Generate report with new format
                            report = extractor.generate_report_from_text(observations_text, st.session_state.user_info)
                            st.session_state.report_generated = report

                            # Get child ID from mappings if available
                            child_id = structured_data.get("studentId", "")
                            if not child_id:
                                # Try to find child by name
                                child_data = supabase.table('users').select("id").ilike("name",
                                                                                        f"%{structured_data.get('studentName', '')}%").execute().data
                                if child_data:
                                    child_id = child_data[0]['id']

                            # Insert observation and capture the returned ID
                            observation_response = supabase.table('observations').insert({
                                "student_id": selected_child_id,
                                "username": st.session_state.auth['user_id'],
                                "student_name": structured_data.get("studentName", ""),
                                "observer_name": st.session_state.user_info['observer_name'],
                                "class_name": structured_data.get("className", ""),
                                "date": structured_data.get("date", ""),
                                "observations": observations_text,
                                "strengths": json.dumps(structured_data.get("strengths", [])),
                                "areas_of_development": json.dumps(structured_data.get("areasOfDevelopment", [])),
                                "recommendations": json.dumps(structured_data.get("recommendations", [])),
                                "timestamp": datetime.now().isoformat(),
                                "filename": uploaded_file.name,
                                "full_data": json.dumps(structured_data),
                                "theme_of_day": structured_data.get("themeOfDay", ""),
                                "curiosity_seed": structured_data.get("curiositySeed", ""),
                                "file_url": file_url  # Add the file URL
                            }).execute()

                            # Get the observation ID for goal alignment
                            observation_id = observation_response.data[0]['id'] if observation_response.data else None

                            st.success("Data processed and saved successfully!")

                            # Analyze alignment with goals if we have a child ID and observation ID
                            if child_id and observation_id:
                                goals = supabase.table('goals').select("*").eq("child_id", child_id).eq("status",
                                                                                                        "active").execute().data

                                if goals:
                                    for goal in goals:
                                        # Use Groq API to analyze alignment
                                        alignment_prompt = f"""
                                        Analyze how well this observation report aligns with the following learning goal:

                                        GOAL: {goal['goal_text']}

                                        OBSERVATION REPORT:
                                        {observations_text}

                                        Provide your analysis in JSON format with:
                                        - alignment_score (0-10 scale)
                                        - analysis_text (detailed explanation of alignment)
                                        - suggested_next_steps
                                        """

                                        try:
                                            # Using Groq API for analysis
                                            response = requests.post(
                                                'https://api.groq.com/openai/v1/chat/completions',
                                                headers={
                                                    'Authorization': f'Bearer {extractor.groq_api_key}',
                                                    'Content-Type': 'application/json'
                                                },
                                                json={
                                                    "model": "llama-3.3-70b-versatile",
                                                    "messages": [
                                                        {
                                                            "role": "system",
                                                            "content": "You are an educational assessment AI that analyzes how well observation reports align with learning goals."
                                                        },
                                                        {
                                                            "role": "user",
                                                            "content": alignment_prompt
                                                        }
                                                    ],
                                                    "temperature": 0.2,
                                                    "response_format": {"type": "json_object"}
                                                }
                                            )

                                            response.raise_for_status()
                                            data = response.json()
                                            analysis = json.loads(data['choices'][0]['message']['content'])

                                            # Save alignment analysis
                                            alignment_data = {
                                                "goal_id": goal['id'],
                                                "report_id": observation_id,  # Use the actual observation ID
                                                "alignment_score": analysis.get('alignment_score', 0),
                                                "analysis_text": analysis.get('analysis_text', 'No analysis')
                                            }
                                            supabase.table('goal_alignments').insert(alignment_data).execute()

                                        except Exception as e:
                                            st.error(f"Goal alignment analysis failed: {str(e)}")
                        else:
                            st.error("No observations found in the extracted data")
                    except Exception as e:
                        st.error(f"Processing error: {str(e)}")


        # Audio Processing
        elif st.session_state.processing_mode == "audio":
            st.info("Audio Mode: Upload an audio recording of an observation session")
            uploaded_file = st.file_uploader("Choose an audio file",
                                             type=["wav", "mp3", "m4a", "mpeg", "ogg", "flac", "aac", "wma", "aiff"])
            if uploaded_file and st.button("Process & Generate Report"):
                if not assemblyai_key:
                    st.error("AssemblyAI API key is missing.")
                else:
                    with st.spinner("Step 1/2: Transcribing audio..."):
                        transcript = extractor.transcribe_with_assemblyai(uploaded_file)
                        st.session_state.audio_transcription = transcript

                    with st.spinner("Step 2/2: Generating report..."):
                        # Upload file to storage
                        file_url = None
                        if uploaded_file:
                            # Reset file pointer to beginning
                            uploaded_file.seek(0)
                            # Upload to storage
                            file_url = upload_file_to_storage(
                                uploaded_file.getvalue(),
                                uploaded_file.name,
                                f"audio/{uploaded_file.type.split('/')[1]}"
                            )

                        # Generate report with new format
                        report = extractor.generate_report_from_text(transcript, st.session_state.user_info)
                        st.session_state.report_generated = report

                        # Extract theme and curiosity seed using AI
                        try:
                            theme_prompt = f"""
                            Based on this transcript, identify:
                            1. The main theme or topic of the day
                            2. A curiosity seed (something that sparked the child's interest or curiosity)

                            Return as JSON with keys: "themeOfDay" and "curiositySeed"

                            TRANSCRIPT:
                            {transcript}
                            """

                            model = genai.GenerativeModel('gemini-2.0-flash-002')
                            response = model.generate_content([
                                {"role": "user", "parts": [{"text": theme_prompt}]}
                            ])

                            theme_data = json.loads(response.text)
                            theme_of_day = theme_data.get("themeOfDay", "")
                            curiosity_seed = theme_data.get("curiositySeed", "")
                        except:
                            theme_of_day = ""
                            curiosity_seed = ""

                        supabase.table('observations').insert({
                            "student_id": selected_child_id,
                            "username": st.session_state.auth['user_id'],
                            "student_name": st.session_state.user_info['student_name'],
                            "observer_name": st.session_state.user_info['observer_name'],
                            "class_name": "",
                            "date": st.session_state.user_info['session_date'],
                            "observations": transcript,
                            "strengths": json.dumps([]),
                            "areas_of_development": json.dumps([]),
                            "recommendations": json.dumps([]),
                            "timestamp": datetime.now().isoformat(),
                            "filename": uploaded_file.name,
                            "full_data": json.dumps({"transcript": transcript, "report": report}),
                            "theme_of_day": theme_of_day,
                            "curiosity_seed": curiosity_seed,
                            "file_url": file_url  # Add the file URL
                        }).execute()

        # Transcript Editor
        if st.session_state.audio_transcription:
            if st.button("Edit Transcription" if not st.session_state.show_edit_transcript else "Hide Editor"):
                st.session_state.show_edit_transcript = not st.session_state.show_edit_transcript
            if st.session_state.show_edit_transcript:
                st.subheader("Edit Transcription")
                edited = st.text_area("Edit transcript below:", value=st.session_state.audio_transcription, height=200)
                if edited != st.session_state.audio_transcription:
                    st.session_state.audio_transcription = edited
                if st.button("Regenerate Report with Edited Transcript"):
                    with st.spinner("Regenerating report..."):
                        report = extractor.generate_report_from_text(edited, st.session_state.user_info)
                        st.session_state.report_generated = report

        # Report Display and Download
        if st.session_state.report_generated:
            st.subheader("Generated Report")
            st.markdown(st.session_state.report_generated)
            docx_file = extractor.create_word_document(st.session_state.report_generated)
            student = st.session_state.user_info['student_name'].replace(" ", "_")
            date = st.session_state.user_info['session_date'].replace("/", "-")
            filename = f"Observer_Report_{student}_{date}.docx"
            st.download_button("Download as Word Document", data=docx_file, file_name=filename,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            st.subheader("Email Report")
            with st.form("email_form"):
                to_email = st.text_input("Recipient Email", value="parent@example.com")
                subject = st.text_input("Subject",
                                        value=f"Observer Report for {st.session_state.user_info['student_name']}")
                submitted = st.form_submit_button("Send Email")
                if submitted:
                    success, message = extractor.send_email(to_email, subject, st.session_state.report_generated)
                    if success:
                        st.success(message)
                    else:
                        st.error(message)

    with observer_tabs[1]:
        st.subheader("Goal Management")

        # Get all children assigned to this observer
        mappings = supabase.table('observer_child_mappings').select("child_id").eq("observer_id", st.session_state.auth[
            'user_id']).execute().data
        child_ids = [m['child_id'] for m in mappings]

        if child_ids:
            # Get child details
            children = supabase.table('children').select("*").in_("id", child_ids).execute().data
            child_options = {c['id']: c.get('name', f"Child {c['id'][:4]}") for c in children}

            # Form to add new goal
            with st.expander("Add New Goal"):
                with st.form("add_goal"):
                    selected_child = st.selectbox("Select Child", options=list(child_options.keys()),
                                                  format_func=lambda x: child_options[x])
                    goal_text = st.text_area("Goal Description", height=100)
                    target_date = st.date_input("Target Date", min_value=datetime.now().date())

                    if st.form_submit_button("Save Goal"):
                        if goal_text.strip():
                            goal_data = {
                                "observer_id": st.session_state.auth['user_id'],
                                "child_id": selected_child,
                                "goal_text": goal_text,
                                "target_date": target_date.isoformat()
                            }
                            supabase.table('goals').insert(goal_data).execute()
                            st.success("Goal saved successfully!")
                        else:
                            st.error("Please enter a goal description")

            # Display current goals
            st.subheader("Current Goals")
            goals = supabase.table('goals').select("*").eq("observer_id",
                                                           st.session_state.auth['user_id']).execute().data

            if goals:
                for goal in goals:
                    child_name = child_options.get(goal['child_id'], "Unknown Child")
                    with st.expander(f"Goal for {child_name} (Due: {goal.get('target_date', 'No date')})"):
                        st.write(goal['goal_text'])
                        st.write(f"Status: {goal.get('status', 'active')}")

                        # Display alignment scores if available
                        alignments = supabase.table('goal_alignments').select("*").eq("goal_id",
                                                                                      goal['id']).execute().data
                        if alignments:
                            st.write("**Alignment with Reports:**")
                            for alignment in alignments:
                                report = supabase.table('observations').select("date").eq("id", alignment[
                                    'report_id']).execute().data
                                report_date = report[0]['date'] if report else "Unknown date"
                                st.write(f"- {report_date}: Score {alignment.get('alignment_score', 0)}/10")
                                if alignment.get('analysis_text'):
                                    with st.expander("Analysis Details"):
                                        st.write(alignment['analysis_text'])

                        # Option to mark as complete/delete
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("Mark Achieved", key=f"complete_{goal['id']}"):
                                supabase.table('goals').update({"status": "achieved"}).eq("id", goal['id']).execute()
                                st.rerun()
                        with col2:
                            if st.button("Delete", key=f"delete_{goal['id']}"):
                                supabase.table('goals').delete().eq("id", goal['id']).execute()
                                st.rerun()
            else:
                st.info("No goals set yet")
        else:
            st.warning("No children assigned to you yet")

    with observer_tabs[2]:  # Messages tab
        st.subheader("Messages with Parents")

        # Get all children assigned to this observer
        mappings = supabase.table('observer_child_mappings').select("child_id").eq("observer_id",
                                                                                   st.session_state.auth[
                                                                                       'user_id']).execute().data
        child_ids = [m['child_id'] for m in mappings]

        if not child_ids:
            st.warning("No children assigned to you yet")
        else:
            # Get all parents of these children
            parents = supabase.table('users').select("*").eq("role", "Parent").in_("child_id", child_ids).execute().data

            if not parents:
                st.info("No parents found for your assigned children")
            else:
                # Create a selection for the parent to message
                parent_options = {p['id']: f"{p.get('name', 'Parent')} - {p.get('email', 'No email')}" for p in parents}
                selected_parent = st.selectbox("Select Parent to Message",
                                               options=list(parent_options.keys()),
                                               format_func=lambda x: parent_options[x])

                # Get child info for this parent
                selected_parent_data = next((p for p in parents if p['id'] == selected_parent), None)
                if selected_parent_data:
                    child_id = selected_parent_data.get('child_id')
                    child_data = supabase.table('children').select("*").eq("id", child_id).execute().data
                    child_name = child_data[0].get('name', 'Child') if child_data else 'Child'

                    st.write(f"**Messaging with parent of:** {child_name}")

                    # Get existing messages
                    messages = supabase.table('messages').select("*") \
                        .or_(
                        f"sender_id.eq.{st.session_state.auth['user_id']},receiver_id.eq.{st.session_state.auth['user_id']}") \
                        .or_(f"sender_id.eq.{selected_parent},receiver_id.eq.{selected_parent}") \
                        .order('timestamp', desc=False) \
                        .execute().data

                    # Display message history
                    message_container = st.container(height=400)
                    with message_container:
                        for msg in messages:
                            is_from_me = msg['sender_id'] == st.session_state.auth['user_id']

                            col1, col2 = st.columns([1, 4])
                            with col1:
                                st.write("You:" if is_from_me else f"{selected_parent_data.get('name', 'Parent')}:")
                            with col2:
                                st.text_area("", value=msg['content'], height=70, key=f"msg_{msg['id']}", disabled=True)

                    # Send new message
                    with st.form("send_message_form_observer"):
                        new_message = st.text_area("Type your message:", height=100)
                        if st.form_submit_button("Send Message"):
                            if new_message.strip():
                                # Insert message into database
                                supabase.table('messages').insert({
                                    "sender_id": st.session_state.auth['user_id'],
                                    "receiver_id": selected_parent,
                                    "content": new_message,
                                    "timestamp": datetime.now().isoformat(),
                                    "read": False
                                }).execute()
                                st.success("Message sent!")
                                st.rerun()
                            else:
                                st.error("Cannot send empty message")

    with observer_tabs[3]:  # Monthly Reports tab
        observer_monthly_report_section(st.session_state.auth['user_id'])

        # Add section to view parent feedback
        st.subheader("Parent Feedback on Monthly Reports")

        # Get all children assigned to this observer
        mappings = supabase.table('observer_child_mappings').select("child_id").eq("observer_id", st.session_state.auth[
            'user_id']).execute().data
        child_ids = [m['child_id'] for m in mappings]

        if child_ids:
            # Get all monthly reports with feedback for these children
            reports = supabase.table('monthly_reports').select("*") \
                .in_("child_id", child_ids) \
                .not_.is_("feedback", "null") \
                .order("year", desc=True) \
                .order("month", desc=True) \
                .execute().data

            if reports:
                for report in reports:
                    child = supabase.table('children').select("name").eq("id", report['child_id']).execute().data
                    child_name = child[0]['name'] if child else "Unknown Child"

                    with st.expander(
                            f"Feedback for {child_name} - {calendar.month_name[report['month']]} {report['year']}"):
                        st.write(f"**Rating:** {'⭐' * report.get('rating', 0)}")
                        st.write(f"**Feedback:** {report['feedback']}")
                        st.write(f"**Submitted on:** {report.get('feedback_submitted_at', 'Unknown date')}")
            else:
                st.info("No feedback received yet")
        else:
            st.warning("No children assigned to you yet")


if __name__ == "__main__":
    main()
